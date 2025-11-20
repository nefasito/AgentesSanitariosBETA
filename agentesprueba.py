import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import io
import locale
# --- IMPORT NECESARIO PARA GENERAR ARCHIVOS .docx ---
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# ---------------------------
# Funciones auxiliares (Demografía)
# ---------------------------

def estimar_pea(edad, genero):
    """
    Estima si una persona pertenece a la Población Económicamente Activa (PEA).
    Se basa en rangos de edad genéricos para hombres y mujeres.
    """
    if pd.isna(edad) or pd.isna(genero):
        return False
    if isinstance(genero, str):
        genero = genero.lower()
    try:
        edad = int(edad)
    except (ValueError, TypeError):
        return False

    if genero == "masculino":
        return 18 <= edad <= 65
    elif genero == "femenino":
        return 18 <= edad <= 60
    else:
        return False

def calcular_indices_avanzados(df_original, df_sin_duplicados):
    """
    Calcula los principales índices demográficos y las nuevas métricas del censo.
    Utiliza el dataframe original para el conteo de viviendas y el de-duplicado para el resto.
    """
    # Métricas basadas en datos únicos (sin DNI duplicados)
    df = df_sin_duplicados
    personas_unicas = len(df)
    
    # Manejar el caso de dataframe vacío
    if personas_unicas == 0:
        return {
            "Personas censadas (únicas)": 0,
            "Viviendas relevadas (estimado)": 0,
            "Población Económicamente Activa (PEA)": 0,
            "Población Económicamente Inactiva (PEI)": 0,
            "Población por grupos de edad": {
                "0-14 años (Joven)": 0,
                "15-64 años (Adulta)": 0,
                "65+ años (Mayor)": 0,
                "80+ años": 0,
            },
            "Índices de dependencia": {
                "Índice de dependencia total": np.nan,
                "Índice de dependencia económica": np.nan,
                "Tasa de dependencia juvenil": np.nan,
                "Tasa de dependencia de la tercera edad": np.nan
            },
            "Índice de envejecimiento": np.nan,
            "Relación por género": {
                "Índice de masculinidad": np.nan,
                "Índice de feminidad": np.nan,
            },
            "Índice de longevidad": np.nan,
            "Edad mediana": np.nan,
            "Edad promedio": np.nan
        }

    menores_15 = df[df["EDAD"] < 15].shape[0]
    poblacion_15_64 = df[(df["EDAD"] >= 15) & (df["EDAD"] <= 64)].shape[0]
    mayores_65 = df[df["EDAD"] >= 65].shape[0]
    poblacion_80_mas = df[df["EDAD"] >= 80].shape[0]
    pea = df[df["PEA_estimado"] == True].shape[0]

    df['GENERO_lower'] = df['GENERO'].str.lower()
    mujeres = df[df['GENERO_lower'] == "femenino"].shape[0]
    hombres = df[df['GENERO_lower'] == "masculino"].shape[0]
    
    edad_mediana = df["EDAD"].median()
    edad_promedio = df["EDAD"].mean()

    # Nuevas métricas (Usa la columna 'DIRECCION' estandarizada)
    viviendas_relevadas = df_original['DIRECCION'].nunique() if 'DIRECCION' in df_original.columns else 0

    return {
        "Personas censadas (únicas)": personas_unicas,
        "Viviendas relevadas (estimado)": viviendas_relevadas,
        "Población Económicamente Activa (PEA)": pea,
        "Población Económicamente Inactiva (PEI)": personas_unicas - pea,
        "Población por grupos de edad": {
            "0-14 años (Joven)": menores_15,
            "15-64 años (Adulta)": poblacion_15_64,
            "65+ años (Mayor)": mayores_65,
            "80+ años": poblacion_80_mas,
        },
        "Índices de dependencia": {
            "Índice de dependencia total": round(((menores_15 + mayores_65) / personas_unicas) * 100, 2) if personas_unicas > 0 else np.nan,
            "Índice de dependencia económica": round(((menores_15 + mayores_65) / pea) * 100, 2) if pea > 0 else np.nan,
            "Tasa de dependencia juvenil": round((menores_15 / poblacion_15_64) * 100, 2) if poblacion_15_64 > 0 else np.nan,
            "Tasa de dependencia de la tercera edad": round((mayores_65 / poblacion_15_64) * 100, 2) if poblacion_15_64 > 0 else np.nan
        },
        "Índice de envejecimiento": round((mayores_65 / menores_15) * 100, 2) if menores_15 > 0 else np.nan,
        "Relación por género": {
            "Índice de masculinidad": round((hombres / mujeres) * 100, 2) if mujeres > 0 else np.nan,
            "Índice de feminidad": round((mujeres / hombres) * 100, 2) if hombres > 0 else np.nan,
        },
        "Índice de longevidad": round((poblacion_80_mas / mayores_65) * 100, 2) if mayores_65 > 0 else np.nan,
        "Edad mediana": edad_mediana,
        "Edad promedio": round(edad_promedio, 2)
    }

def calcular_tasas(indices, natalidad_bruta, mortalidad_bruta, tfg, tasa_migracion_neta):
    """
    Calcula tasas demográficas basadas en los índices y en datos simulados por el usuario.
    """
    return {
        "Tasa de natalidad (por 1000)": natalidad_bruta,
        "Tasa de mortalidad (por 1000)": mortalidad_bruta,
        "Tasa de migración neta (por 1000)": tasa_migracion_neta,
        "Tasa de crecimiento natural (%)": round(((natalidad_bruta - mortalidad_bruta) / 10), 2),
        "Tasa de crecimiento anual (%)": round(((natalidad_bruta - mortalidad_bruta + tasa_migracion_neta) / 10), 2),
        "Tasa global de fecundidad (TFG)": tfg,
    }

def interpretar_piramide(indices):
    """
    Interpreta la forma de la pirámide poblacional.
    """
    joven = indices["Población por grupos de edad"]["0-14 años (Joven)"]
    adulta = indices["Población por grupos de edad"]["15-64 años (Adulta)"]
    mayor = indices["Población por grupos de edad"]["65+ años (Mayor)"]
    total = joven + adulta + mayor

    if total == 0:
        return "Datos insuficientes", "No hay datos para interpretar la pirámide."

    if joven / total > 0.35 and mayor / total < 0.10:
        return "Expansiva (Progresiva)", "La base ancha indica una alta proporción de población joven y altas tasas de natalidad. Sugiere un crecimiento poblacional futuro y es característica de poblaciones jóvenes."
    elif joven / total < 0.25 and mayor / total > 0.15:
        return "Constrictiva (Regresiva)", "La base estrecha y el ensanchamiento en la parte superior indican una baja natalidad y una alta proporción de población adulta y mayor. Característica de poblaciones envejecidas con posible decrecimiento futuro."
    else:
        return "Estacionaria", "La estructura es relatively uniforme en los diferentes grupos de edad, lo que sugiere tasas de natalidad y mortalidad bajas y estables. Indica un crecimiento lento o nulo de la población."

def generar_proyeccion(indices, tasas):
    """
    Genera una proyección simple de la población.
    """
    poblacion_actual = indices["Personas censadas (únicas)"]
    if poblacion_actual == 0:
        return {"5 años": 0, "10 años": 0, "20 años": 0}
        
    # Convertir tasa de % a decimal (ej: 2.1% -> 0.021)
    tasa_crecimiento = tasas["Tasa de crecimiento anual (%)"] / 100
    
    proyeccion = {
        # Fórmula de interés compuesto: P_futura = P_actual * (1 + tasa)^n
        "5 años": int(poblacion_actual * (1 + tasa_crecimiento) ** 5),
        "10 años": int(poblacion_actual * (1 + tasa_crecimiento) ** 10),
        "20 años": int(poblacion_actual * (1 + tasa_crecimiento) ** 20),
    }
    return proyeccion

# --------------------------------
# Funciones de Gráficos (Demografía)
# --------------------------------

def graficar_piramide(df, localidad=""):
    """
    Genera y devuelve una pirámide poblacional con grupos quinquenales y colores pastel,
    y los datos subyacentes.
    """
    if df.empty or df['EDAD'].isnull().all():
        fig, ax = plt.subplots()
        ax.text(0.5, 0.5, "No hay datos para la pirámide", horizontalalignment='center', verticalalignment='center')
        return fig, pd.DataFrame(columns=["Grupo de Edad", "Hombres", "Mujeres"])

    N_total = len(df)
    
    bins = list(range(0, 100, 5)) + [np.inf]
    labels = [f'{i}-{i+4}' for i in range(0, 95, 5)] + ['95+']
    df["Grupo_Edad"] = pd.cut(df["EDAD"], bins=bins, labels=labels, right=False)

    hombres_counts = df[df["GENERO"].str.lower() == "masculino"]["Grupo_Edad"].value_counts().sort_index()
    mujeres_counts = df[df["GENERO"].str.lower() == "femenino"]["Grupo_Edad"].value_counts().sort_index()

    all_groups = pd.CategoricalIndex(labels, categories=labels, ordered=True)
    hombres_counts = hombres_counts.reindex(all_groups, fill_value=0)
    mujeres_counts = mujeres_counts.reindex(all_groups, fill_value=0)

    grupos = [str(g) for g in hombres_counts.index]

    fig, ax = plt.subplots(figsize=(12, 8))
    
    color_hombres = "#A7D8A9"
    color_mujeres = "#FFB6C1"

    bars_h = ax.barh(grupos, -hombres_counts.values, color=color_hombres, label="Hombres")
    bars_m = ax.barh(grupos, mujeres_counts.values, color=color_mujeres, label="Mujeres")

    ax.bar_label(bars_h, labels=[f'{abs(int(v))}' for v in hombres_counts.values], padding=3, color='black', fontsize=10)
    ax.bar_label(bars_m, labels=[f'{int(v)}' for v in mujeres_counts.values], padding=3, color='black', fontsize=10)
    
    ax.set_xlabel("Población", fontsize=12)
    ax.set_ylabel("Grupos de edad", fontsize=12)
    titulo = f"Pirámide Poblacional de {localidad.title() if localidad else 'la Localidad'} (N={N_total})"
    ax.set_title(titulo, fontsize=18, fontweight='bold')
    ax.legend()
    ax.spines['right'].set_visible(False)
    ax.spines['top'].set_visible(False)
    
    ax.xaxis.set_major_formatter(plt.FuncFormatter(lambda x, loc: str(abs(int(x)))))
    
    if not hombres_counts.empty and not mujeres_counts.empty:
        max_val = max(hombres_counts.max(), mujeres_counts.max()) if hombres_counts.max() > 0 or mujeres_counts.max() > 0 else 1
        ax.set_xlim(-max_val * 1.25, max_val * 1.25)
    
    fig.tight_layout()

    # Preparar datos para la tabla
    df_piramide_data = pd.DataFrame({
        "Grupo de Edad": grupos,
        "Hombres": hombres_counts.values,
        "Mujeres": mujeres_counts.values
    })
    
    return fig, df_piramide_data

def graficar_genero(df, localidad=""):
    if df.empty or df['GENERO'].isnull().all():
        fig, ax = plt.subplots()
        ax.text(0.5, 0.5, "No hay datos de género", horizontalalignment='center', verticalalignment='center')
        return fig

    conteo_genero = df["GENERO"].str.lower().value_counts()
    
    if conteo_genero.empty:
        fig, ax = plt.subplots()
        ax.text(0.5, 0.5, "No hay datos de género", horizontalalignment='center', verticalalignment='center')
        return fig
    
    N_total = conteo_genero.sum()
        
    color_map = {"masculino": "#A7D8A9", "femenino": "#FFB6C1"}
    labels = conteo_genero.index.str.capitalize()
    colors = [color_map.get(g.lower(), "#CCCCCC") for g in conteo_genero.index]
    
    fig, ax = plt.subplots(figsize=(8, 8))
    
    def autopct_format(p):
        count = int(round(p * N_total / 100.0))
        return f'{p:.1f}%\n({count})'

    ax.pie(conteo_genero.values, labels=labels, autopct=autopct_format, startangle=90, colors=colors)
    ax.axis('equal')
    
    titulo = f"Distribución por Género en {localidad.title() if localidad else 'la Localidad'} (N={N_total})"
    ax.set_title(titulo, fontsize=18, fontweight='bold')
    fig.tight_layout()
    return fig

def graficar_analisis_adicional(df):
    """
    Genera gráficos para variables socio-económicas (NO de salud).
    """
    figs = {}
    
    if df.empty:
        return figs

    # Gráfico de Situación Laboral
    if 'SITUACION_LABORAL' in df.columns:
        counts = df['SITUACION_LABORAL'].value_counts()
        if not counts.empty:
            N_laboral = counts.sum()
            fig, ax = plt.subplots(figsize=(10, 6))
            bars = counts.plot(kind='bar', ax=ax, color='#A7D8A9')
            
            ax.set_title(f'Distribución por Situación Laboral (N={N_laboral})', fontsize=16, fontweight='bold')
            ax.set_ylabel('Cantidad de Personas', fontsize=12)
            ax.tick_params(axis='x', rotation=45)
            ax.spines['right'].set_visible(False)
            ax.spines['top'].set_visible(False)
            
            if ax.containers:
                ax.bar_label(ax.containers[0], fmt='%.0f', fontsize=10, padding=3)
            
            fig.tight_layout()
            figs['situacion_laboral'] = fig

    # Gráfico de Tenencia de Vivienda
    if 'VIVIENDA_TENENCIA' in df.columns:
        counts = df['VIVIENDA_TENENCIA'].value_counts()
        if not counts.empty:
            N_vivienda = counts.sum()
            fig, ax = plt.subplots(figsize=(10, 6))
            bars = counts.plot(kind='bar', ax=ax, color='#FFB6C1')

            ax.set_title(f'Distribución por Tenencia de Vivienda (N={N_vivienda})', fontsize=16, fontweight='bold')
            ax.set_ylabel('Cantidad de Personas', fontsize=12)
            ax.tick_params(axis='x', rotation=45)
            ax.spines['right'].set_visible(False)
            ax.spines['top'].set_visible(False)

            if ax.containers:
                ax.bar_label(ax.containers[0], fmt='%.0f', fontsize=10, padding=3)

            fig.tight_layout()
            figs['vivienda_tenencia'] = fig
            
    return figs

# --------------------------------
# NUEVAS Funciones (Salud)
# --------------------------------

def calcular_estadisticas_salud(df):
    """
    Calcula estadísticas para enfermedades, control de niño, vacunas y cobertura de obra social.
    """
    TODAY = pd.to_datetime('today').normalize() 
        
    stats = {
        'enfermedades': {},
        'control_nino': {},
        'vacunas': {},
        'enfermedades_obra_social': {},
        'obra_social_general': {},
        'embarazadas': {},
        'mayores_65_os': {},
        'menores_15_os': {}
    }
    
    filtros_booleanos = {}
    
    if df.empty:
        return stats, filtros_booleanos 

    # --- Procesamiento de OBRA_SOCIAL ---
    if 'OBRA_SOCIAL' in df.columns:
        df['OBRA_SOCIAL_limpia'] = df['OBRA_SOCIAL'].astype(str).str.strip().str.replace('sí', 'si', case=False).str.lower()
        
        total_con_os = (df['OBRA_SOCIAL_limpia'] == 'si').sum()
        total_sin_os = (df['OBRA_SOCIAL_limpia'] == 'no').sum()
        total_poblacion = len(df)
        
        indice_general = (total_con_os / total_poblacion) * 100 if total_poblacion > 0 else 0
        
        stats['obra_social_general'] = {
            'Con Obra Social': total_con_os,
            'Sin Obra Social': total_sin_os,
            'Población Total (única)': total_poblacion,
            'Indice Mutualizacion (%)': round(indice_general, 2)
        }
    else:
        df['OBRA_SOCIAL_limpia'] = 'sin datos'

    # 1. Estadísticas de Enfermedades
    ENFERMEDADES_COLS = ['DIABETES', 'HTA', 'CHAGAS', 'TBC', 'CANCER']
    enfermedades_presentes = [c for c in ENFERMEDADES_COLS if c in df.columns]
    
    filtros_enfermedades_base = {}
    for col in enfermedades_presentes:
        valor_limpio = df[col].astype(str).str.strip().str.replace('sí', 'si', case=False).str.lower()
        filtro_bool = (valor_limpio == 'si')
        filtros_enfermedades_base[col] = filtro_bool
        filtros_booleanos[col] = filtro_bool

    # 2. Añadir el filtro DUAL si existen DIABETES y HTA
    if 'DIABETES' in filtros_enfermedades_base and 'HTA' in filtros_enfermedades_base:
        filtro_dual = filtros_enfermedades_base['DIABETES'] & filtros_enfermedades_base['HTA']
        filtros_enfermedades_base['DUAL (DBT+HTA)'] = filtro_dual
        filtros_booleanos['DUAL (DBT+HTA)'] = filtro_dual


    # 3. Procesar todos los filtros
    counts_enf = {}
    if filtros_enfermedades_base:
        for col, filtro_enfermedad in filtros_enfermedades_base.items():
            
            counts_enf[col] = filtro_enfermedad.sum()

            # --- Cálculo de OS para esta enfermedad ---
            if 'OBRA_SOCIAL' in df.columns:
                df_enfermos = df[filtro_enfermedad]
                
                con_os = (df_enfermos['OBRA_SOCIAL_limpia'] == 'si').sum()
                sin_os = (df_enfermos['OBRA_SOCIAL_limpia'] == 'no').sum()
                total_enf = filtro_enfermedad.sum()

                con_os_pct = (con_os / total_enf) * 100 if total_enf > 0 else 0
                sin_os_pct = (sin_os / total_enf) * 100 if total_enf > 0 else 0
                
                stats['enfermedades_obra_social'][col] = {
                    'Total Enfermos': total_enf,
                    'Con Obra Social': con_os,
                    'Sin Obra Social': sin_os,
                    'Con OS (%)': round(con_os_pct, 2),
                    'Sin OS (%)': round(sin_os_pct, 2)
                }
    
    if counts_enf: 
        if sum(counts_enf.values()) > 0:
            stats['enfermedades'] = counts_enf

    # 2. Estadísticas de Control Niño
    CONTROL_NINO_COL = 'NIÑO_CONTROLES_COMPLETOS' 
    if CONTROL_NINO_COL in df.columns:
        valor_limpio_cn = df[CONTROL_NINO_COL].astype(str).str.strip().str.replace('sí', 'si', case=False).str.lower()
        mapa_control = {
            'si': 'Completo', 'completo': 'Completo',
            'no': 'Incompleto', 'incompleto': 'Incompleto'
        }
        mapped_cn = valor_limpio_cn.map(mapa_control)
        counts_cn = mapped_cn.dropna().value_counts()
        stats['control_nino'] = counts_cn.to_dict()

    # 3. Estadísticas de Vacunas
    VACUNAS_COL = 'NIÑO_VACUNACION_COMPLETA'
    if VACUNAS_COL in df.columns:
        valor_limpio_vac = df[VACUNAS_COL].astype(str).str.strip().str.replace('sí', 'si', case=False).str.lower()
        mapa_vacunas = {
            'si': 'Completas', 'completo': 'Completas', 'completas': 'Completas',
            'no': 'Incompletas', 'incompleto': 'Incompletas', 'incompletas': 'Incompletas'
        }
        mapped_vac = valor_limpio_vac.map(mapa_vacunas)
        counts_vac = mapped_vac.dropna().value_counts()
        stats['vacunas'] = counts_vac.to_dict()

    # 4. Estadísticas de Embarazo
    EMBARAZO_FUM_COL = 'MUJER_FUM'
    if EMBARAZO_FUM_COL in df.columns:
        fum_dates = pd.to_datetime(df[EMBARAZO_FUM_COL], errors='coerce', dayfirst=True)
        approx_fpp = pd.Series(pd.NaT, index=df.index, dtype='datetime64[ns]')
        if not fum_dates.isna().all():
            try:
                approx_fpp = fum_dates + pd.DateOffset(months=9)
            except Exception:
                approx_fpp = fum_dates + pd.to_timedelta(270, unit='D') 

        filtro_embarazadas = (approx_fpp >= TODAY) & (fum_dates.notna())
        total_embarazadas = filtro_embarazadas.sum()
        
        emb_con_os = 0
        emb_sin_os = 0
        emb_con_os_pct = 0
        emb_sin_os_pct = 0

        if total_embarazadas > 0 and 'OBRA_SOCIAL_limpia' in df.columns:
            df_embarazadas = df[filtro_embarazadas]
            emb_con_os = (df_embarazadas['OBRA_SOCIAL_limpia'] == 'si').sum()
            emb_sin_os = (df_embarazadas['OBRA_SOCIAL_limpia'] == 'no').sum()
            emb_con_os_pct = (emb_con_os / total_embarazadas) * 100
            emb_sin_os_pct = (emb_sin_os / total_embarazadas) * 100

        stats['embarazadas'] = {
            'Total Estimado': total_embarazadas,
            'Con Obra Social': emb_con_os,
            'Sin Obra Social': emb_sin_os,
            'Con OS (%)': round(emb_con_os_pct, 2),
            'Sin OS (%)': round(emb_sin_os_pct, 2)
        }
        
        filtros_booleanos['EMBARAZADAS'] = filtro_embarazadas
        

    # 5. Cálculo OS por Grupos de Edad
    if 'OBRA_SOCIAL_limpia' in df.columns and 'EDAD' in df.columns:
        
        # 5.a. Mayores de 65 años
        filtro_mayores_65 = df['EDAD'] >= 65
        df_mayores_65 = df[filtro_mayores_65]
        total_mayores_65 = len(df_mayores_65)
        
        mayores_con_os = (df_mayores_65['OBRA_SOCIAL_limpia'] == 'si').sum()
        mayores_sin_os = (df_mayores_65['OBRA_SOCIAL_limpia'] == 'no').sum()
        mayores_con_os_pct = (mayores_con_os / total_mayores_65) * 100 if total_mayores_65 > 0 else 0
        mayores_sin_os_pct = (mayores_sin_os / total_mayores_65) * 100 if total_mayores_65 > 0 else 0

        stats['mayores_65_os'] = {
            'Total Mayores 65': total_mayores_65,
            'Con Obra Social': mayores_con_os,
            'Sin Obra Social': mayores_sin_os,
            'Con OS (%)': round(mayores_con_os_pct, 2),
            'Sin OS (%)': round(mayores_sin_os_pct, 2)
        }

        # 5.b. Menores de 15 años (Grupo 0-14)
        filtro_menores_15 = df['EDAD'] < 15 # Corresponde a 0-14
        df_menores_15 = df[filtro_menores_15]
        total_menores_15 = len(df_menores_15)

        menores_con_os = (df_menores_15['OBRA_SOCIAL_limpia'] == 'si').sum()
        menores_sin_os = (df_menores_15['OBRA_SOCIAL_limpia'] == 'no').sum()
        menores_con_os_pct = (menores_con_os / total_menores_15) * 100 if total_menores_15 > 0 else 0
        menores_sin_os_pct = (menores_sin_os / total_menores_15) * 100 if total_menores_15 > 0 else 0

        stats['menores_15_os'] = {
            'Total Menores 15 (0-14)': total_menores_15,
            'Con Obra Social': menores_con_os,
            'Sin Obra Social': menores_sin_os,
            'Con OS (%)': round(menores_con_os_pct, 2),
            'Sin OS (%)': round(menores_sin_os_pct, 2)
        }

    return stats, filtros_booleanos

def graficar_estadisticas_salud(stats_dict, localidad="", N_total_poblacion=0):
    """
    Genera gráficos de barras para las estadísticas de salud calculadas.
    """
    figs = {}
    
    # Gráfico de Enfermedades
    if stats_dict['enfermedades']:
        fig_enf, ax_enf = plt.subplots(figsize=(10, 6))
        data_enf = pd.Series(stats_dict['enfermedades']).sort_values(ascending=False)
        bars = data_enf.plot(kind='bar', ax=ax_enf, color='#A7D8A9')
        
        ax_enf.set_title(f'Prevalencia de Enfermedades en {localidad.title()} (N={N_total_poblacion})', fontsize=16, fontweight='bold')
        ax_enf.set_ylabel('Cantidad de Casos Positivos', fontsize=12)
        ax_enf.tick_params(axis='x', rotation=45)
        ax_enf.spines['right'].set_visible(False)
        ax_enf.spines['top'].set_visible(False)
        
        if ax_enf.containers:
            ax_enf.bar_label(ax_enf.containers[0], fmt='%.0f', fontsize=10, padding=3)
        
        fig_enf.tight_layout()
        figs['enfermedades'] = fig_enf
    else:
        figs['enfermedades'] = None

    # Gráfico de Control Niño
    if stats_dict['control_nino']:
        fig_cn, ax_cn = plt.subplots(figsize=(8, 5))
        data_cn = pd.Series(stats_dict['control_nino']).sort_index() 
        N_cn = data_cn.sum()
        
        bars = data_cn.plot(kind='bar', ax=ax_cn, color=['#A7D8A9', '#FFB6C1'])
        
        ax_cn.set_title(f'Estado de Control Niño en {localidad.title()} (N={N_cn})', fontsize=16, fontweight='bold')
        ax_cn.set_ylabel('Cantidad de Niños', fontsize=12)
        ax_cn.tick_params(axis='x', rotation=0)
        ax_cn.spines['right'].set_visible(False)
        ax_cn.spines['top'].set_visible(False)
        
        if ax_cn.containers:
            ax_cn.bar_label(ax_cn.containers[0], fmt='%.0f', fontsize=10, padding=3)

        fig_cn.tight_layout()
        figs['control_nino'] = fig_cn
    else:
        figs['control_nino'] = None

    # Gráfico de Vacunas
    if stats_dict['vacunas']:
        fig_vac, ax_vac = plt.subplots(figsize=(8, 5))
        data_vac = pd.Series(stats_dict['vacunas']).sort_index() 
        N_vac = data_vac.sum()
        
        bars = data_vac.plot(kind='bar', ax=ax_vac, color=['#A7D8A9', '#FFB6C1'])
        
        ax_vac.set_title(f'Estado de Vacunación en {localidad.title()} (N={N_vac})', fontsize=16, fontweight='bold')
        ax_vac.set_ylabel('Cantidad de Personas', fontsize=12)
        ax_vac.tick_params(axis='x', rotation=0)
        ax_vac.spines['right'].set_visible(False)
        ax_vac.spines['top'].set_visible(False)

        if ax_vac.containers:
            ax_vac.bar_label(ax_vac.containers[0], fmt='%.0f', fontsize=10, padding=3)

        fig_vac.tight_layout()
        figs['vacunas'] = fig_vac
    else:
        figs['vacunas'] = None

    return figs


# ---------------------------
# NUEVA Función - Explicación de Fórmulas (Streamlit)
# ---------------------------
def mostrar_explicacion_formulas():
    """
    Muestra un expander en Streamlit con la explicación de todas las fórmulas utilizadas.
    """
    with st.expander("📖 Ver Explicación de Fórmulas y Métricas"):
        st.info("Todas las métricas se calculan utilizando los datos únicos por DNI, excepto las 'Viviendas relevadas'.")

        st.subheader("1. Índices Demográficos")
        st.markdown("**Tasa de dependencia juvenil (%)**")
        st.latex(r"""
        \left( \frac{\text{Población de 0-14 años}}{\text{Población de 15-64 años}} \right) \times 100
        """)
        st.markdown("**Tasa de dependencia de la tercera edad (%)**")
        st.latex(r"""
        \left( \frac{\text{Población de 65+ años}}{\text{Población de 15-64 años}} \right) \times 100
        """)
        st.markdown("**Índice de dependencia total (%)**")
        st.latex(r"""
        \left( \frac{\text{Población 0-14 años} + \text{Población 65+ años}}{\text{Población Total (única)}} \right) \times 100
        """)
        st.markdown("**Índice de dependencia económica (%)**")
        st.latex(r"""
        \left( \frac{\text{Población 0-14 años} + \text{Población 65+ años}}{\text{Población Económicamente Activa (PEA)}} \right) \times 100
        """)
        st.markdown("**Índice de envejecimiento (%)**")
        st.latex(r"""
        \left( \frac{\text{Población de 65+ años}}{\text{Población de 0-14 años}} \right) \times 100
        """)
        st.markdown("**Índice de longevidad (%)**")
        st.latex(r"""
        \left( \frac{\text{Población de 80+ años}}{\text{Población de 65+ años}} \right) \times 100
        """)
        st.markdown("**Índice de masculinidad (%)**")
        st.latex(r"""
        \left( \frac{\text{Total Hombres}}{\text{Total Mujeres}} \right) \times 100
        """)

        st.subheader("2. Tasas de Crecimiento (Simuladas)")
        st.markdown("**Tasa de crecimiento natural (%)**")
        st.latex(r"""
        \frac{\text{Tasa Natalidad (x1000)} - \text{Tasa Mortalidad (x1000)}}{10}
        """)
        st.markdown("**Tasa de crecimiento anual (%)**")
        st.latex(r"""
        \frac{\text{Tasa Natalidad (x1000)} - \text{Tasa Mortalidad (x1000)} + \text{Tasa Migración (x1000)}}{10}
        """)
        st.markdown("**Proyección Poblacional**")
        st.latex(r"""
        P_{\text{futura}} = P_{\text{actual}} \times (1 + \text{Tasa Crec. Anual})^{n}
        """)
        st.markdown("Donde `n` es el número de años (5, 10 o 20) y la Tasa de Crecimiento Anual se usa como decimal (ej: 1.5% = 0.015).")

        st.subheader("3. Índices Socio-Sanitarios")
        st.markdown("**Prevalencia de Enfermedad (%)**")
        st.latex(r"""
        \left( \frac{\text{Total Casos Positivos de la Enfermedad}}{\text{Total Personas Censadas (únicas)}} \right) \times 100
        """)
        st.markdown("**Prevalencia DUAL (DBT+HTA)**")
        st.markdown("Se calcula como un 'Y' lógico: `(Persona tiene DIABETES) Y (Persona tiene HTA)`. La prevalencia se calcula sobre el total de personas únicas.")
        
        st.markdown("**Estimación de Embarazadas**")
        st.markdown("Se cuentan las personas donde la columna `MUJER_FUM` (Fecha de Última Menstruación) cumple la siguiente condición:")
        st.latex(r"""
        (\text{MUJER_FUM} + 9 \text{ meses}) \geq \text{Hoy}
        """)
        st.markdown("*(Donde 'Hoy' se define como la fecha actual en que se ejecuta el informe)*")

        st.markdown("**Índice de Mutualización General (%)**")
        st.latex(r"""
        \left( \frac{\text{Total Personas con Obra Social}}{\text{Total Personas Censadas (únicas)}} \right) \times 100
        """)
        st.markdown("**Cobertura OS en Enfermos (%)**")
        st.latex(r"""
        \left( \frac{\text{Pacientes con [Enfermedad] y CON Obra Social}}{\text{Total de Pacientes con [Enfermedad]}} \right) \times 100
        """)
        st.markdown("**Sin Cobertura OS en Enfermos (%)**")
        st.latex(r"""
        \left( \frac{\text{Pacientes con [Enfermedad] y SIN Obra Social}}{\text{Total de Pacientes con [Enfermedad]}} \right) \times 100
        """)
        
# ---------------------------
# NUEVA Función - Apéndice de Fórmulas (Word)
# ---------------------------
def agregar_seccion_formulas(document):
    """
    Añade una sección de apéndice al documento de Word con la explicación de las fórmulas.
    """
    try:
        document.add_page_break()
        document.add_paragraph('Apéndice: Explicación de Fórmulas y Métricas', style='Titulo1')
        
        p_info = document.add_paragraph()
        p_info.add_run("Nota: Todas las métricas se calculan utilizando los datos únicos por DNI, excepto las 'Viviendas relevadas'.").italic = True

        document.add_paragraph('1. Índices Demográficos', style='Titulo2')
        document.add_paragraph('Tasa de dependencia juvenil (%):', style='List Bullet').add_run(" (Población 0-14 / Población 15-64) * 100")
        document.add_paragraph('Tasa de dependencia de la tercera edad (%):', style='List Bullet').add_run(" (Población 65+ / Población 15-64) * 100")
        document.add_paragraph('Índice de dependencia total (%):', style='List Bullet').add_run(" ((Población 0-14 + Población 65+) / Población Total Única) * 100")
        document.add_paragraph('Índice de dependencia económica (%):', style='List Bullet').add_run(" ((Población 0-14 + Población 65+) / PEA) * 100")
        document.add_paragraph('Índice de envejecimiento (%):', style='List Bullet').add_run(" (Población 65+ / Población 0-14) * 100")
        document.add_paragraph('Índice de longevidad (%):', style='List Bullet').add_run(" (Población 80+ / Población 65+) * 100")
        document.add_paragraph('Índice de masculinidad (%):', style='List Bullet').add_run(" (Total Hombres / Total Mujeres) * 100")

        document.add_paragraph('2. Tasas de Crecimiento (Simuladas)', style='Titulo2')
        document.add_paragraph('Tasa de crecimiento natural (%):', style='List Bullet').add_run(" ((Tasa Natalidad x1000 - Tasa Mortalidad x1000) / 10)")
        document.add_paragraph('Tasa de crecimiento anual (%):', style='List Bullet').add_run(" ((Tasa Natalidad x1000 - Tasa Mortalidad x1000 + Tasa Migración x1000) / 10)")
        p_proy = document.add_paragraph('Proyección Poblacional:', style='List Bullet')
        p_proy.add_run(" P_futura = P_actual * (1 + Tasa Crec. Anual)^n")
        p_proy.add_run("\n (Donde 'n' es el n° de años y la Tasa Crec. Anual se usa como decimal, ej: 1.5% = 0.015)")

        document.add_paragraph('3. Índices Socio-Sanitarios', style='Titulo2')
        document.add_paragraph('Prevalencia de Enfermedad (%):', style='List Bullet').add_run(" (Total Casos Positivos / Total Personas Únicas) * 100")
        document.add_paragraph('Prevalencia DUAL (DBT+HTA):', style='List Bullet').add_run(" Personas con (DIABETES='SI' Y HTA='SI')")
        
        document.add_paragraph('Estimación de Embarazadas:', style='List Bullet').add_run(" (MUJER_FUM + 9 meses) >= Hoy (Fecha Actual)") 

        document.add_paragraph('Índice de Mutualización General (%):', style='List Bullet').add_run(" (Total Personas con OS / Total Personas Únicas) * 100")
        document.add_paragraph('Cobertura OS en Enfermos (%):', style='List Bullet').add_run(" (Pacientes [Enfermedad] con OS / Total Pacientes [Enfermedad]) * 100")
        document.add_paragraph('Sin Cobertura OS en Enfermos (%):', style='List Bullet').add_run(" (Pacientes [Enfermedad] sin OS / Total Pacientes [Enfermedad]) * 100")

    except Exception as e:
        print(f"Error al agregar sección de fórmulas al Word: {e}")
        pass

# ---------------------------
# Función de Informe Word
# ---------------------------

def generar_informe_word(indices, tasas, proyecciones, interpretacion, nombre_informe, figs, df_piramide_data, anio="", tasas_prevalencia_enf=None, stats_os_general=None, stats_enf_os=None, stats_embarazadas=None, stats_mayores_65_os=None, stats_menores_15_os=None):
    document = Document()
    
    # --- Configuración de Estilos ---
    styles = document.styles
    style = styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.5

    style_h1 = styles.add_style('Titulo1', WD_STYLE_TYPE.PARAGRAPH)
    font = style_h1.font
    font.name = 'Arial'
    font.size = Pt(16)
    font.bold = True
    
    style_h2 = styles.add_style('Titulo2', WD_STYLE_TYPE.PARAGRAPH)
    font = style_h2.font
    font.name = 'Arial'
    font.size = Pt(14)
    font.bold = True
    
    # --- Configuración de Idioma (Locale) ---
    try:
        locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
    except locale.Error:
        try:
            locale.setlocale(locale.LC_TIME, 'es_ES')
        except locale.Error:
            pass 

    # --- Título y Fecha ---
    titulo_principal = document.add_paragraph()
    titulo_principal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_texto = f"INFORME DEMOGRÁFICO Y SOCIO-SANITARIO\n{nombre_informe.upper()}"
    if anio and anio != "Todos":
        titulo_texto += f"\n(AÑO: {anio})"
    runner = titulo_principal.add_run(titulo_texto)
    runner.bold = True
    runner.font.size = Pt(18)
    
    fecha = document.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha.add_run(f"Fecha de Generación: {pd.Timestamp.now().strftime('%d de %B de %Y')}")

    # --- AVISO IMPORTANTE (Word) ---
    document.add_paragraph() 
    aviso_p = document.add_paragraph()
    aviso_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aviso_run = aviso_p.add_run(
        "AVISO IMPORTANTE\n\n"
        "El presente informe es un borrador generado automáticamente por un "
        "programa en fase de prueba. Los datos y estadísticas aquí presentados pueden "
        "contener errores, incoherencias o inconsistencias.\n\n"
        "Se ruega corroborar toda la información antes de su presentación oficial o difusión.\n"
        "Diseñado por el Tec. Muñoz, Mauricio Iván."
    )
    aviso_run.bold = True
    aviso_run.font.size = Pt(10)
    aviso_run.font.italic = True
    # --- FIN AVISO ---

    document.add_page_break()

    # --- Sección 1: Análisis Demográfico ---
    document.add_paragraph(f'Análisis Demográfico: {nombre_informe.title()}', style='Titulo1')
    document.add_paragraph(
        f"Este informe presenta un análisis detallado de la estructura demográfica de {nombre_informe.title()} "
        f"para el período '{anio}'. Se calculan índices clave, tasas de crecimiento y proyecciones futuras."
    )

    document.add_paragraph('Resultados del Censo', style='Titulo2')
    p = document.add_paragraph()
    p.add_run('Métricas del Censo').bold = True
    document.add_paragraph(f"Personas censadas (únicas): ", style='List Bullet').add_run(str(indices['Personas censadas (únicas)'])).bold = True
    document.add_paragraph(f"Viviendas relevadas (estimado): ", style='List Bullet').add_run(str(indices['Viviendas relevadas (estimado)'])).bold = True
    document.add_paragraph(f"Población Económicamente Activa (PEA): ", style='List Bullet').add_run(str(indices['Población Económicamente Activa (PEA)'])).bold = True
    document.add_paragraph(f"Edad Mediana: ", style='List Bullet').add_run(f"{indices['Edad mediana']} años").bold = True
    
    if stats_os_general:
        p_os = document.add_paragraph()
        p_os.add_run('Métricas de Cobertura').bold = True
        p_os_val = stats_os_general.get('Indice Mutualizacion (%)', 0)
        p_os_total = stats_os_general.get('Con Obra Social', 0)
        document.add_paragraph(f"Índice de Mutualización General: ", style='List Bullet').add_run(f"{p_os_val} %").bold = True
        document.add_paragraph(f"Total con Obra Social: ", style='List Bullet').add_run(f"{p_os_total}").bold = True

    if stats_embarazadas:
        total_emb = stats_embarazadas.get('Total Estimado', 0)
        con_os_emb = stats_embarazadas.get('Con Obra Social', 'N/A')
        sin_os_emb = stats_embarazadas.get('Sin Obra Social', 'N/A')
        
        p_emb = document.add_paragraph(f"Embarazadas (Estimado): ", style='List Bullet')
        p_emb.add_run(f"{total_emb}").bold = True
        p_emb.add_run(f" (Con OS: {con_os_emb} | Sin OS: {sin_os_emb})")

        p_emb_note = document.add_paragraph(
            f"(Cálculo basado en (MUJER_FUM + 9 meses >= Hoy))" 
        )
        p_emb_note.runs[0].font.size = Pt(9)
        p_emb_note.runs[0].italic = True

    if stats_mayores_65_os:
        total_may = stats_mayores_65_os.get('Total Mayores 65', 0)
        con_os_may = stats_mayores_65_os.get('Con Obra Social', 'N/A')
        sin_os_may = stats_mayores_65_os.get('Sin Obra Social', 'N/A')
        p_may = document.add_paragraph(f"Población 65+ años: ", style='List Bullet')
        p_may.add_run(f"{total_may}").bold = True
        p_may.add_run(f" (Con OS: {con_os_may} | Sin OS: {sin_os_may})")

    if stats_menores_15_os:
        total_men = stats_menores_15_os.get('Total Menores 15 (0-14)', 0)
        con_os_men = stats_menores_15_os.get('Con Obra Social', 'N/A')
        sin_os_men = stats_menores_15_os.get('Sin Obra Social', 'N/A')
        p_men = document.add_paragraph(f"Población 0-14 años: ", style='List Bullet')
        p_men.add_run(f"{total_men}").bold = True
        p_men.add_run(f" (Con OS: {con_os_men} | Sin OS: {sin_os_men})")

    p = document.add_paragraph()
    p.add_run('Índices de Dependencia y Envejecimiento').bold = True
    for key, value in indices["Índices de dependencia"].items():
        document.add_paragraph(f"{key}: ", style='List Bullet').add_run(f"{value} %").bold = True
    
    # --- Sección 2: Visualizaciones ---
    document.add_paragraph('Visualizaciones Gráficas', style='Titulo1')
    
    orden_graficos = [
        'piramide', 'genero', 
        'enfermedades', 'control_nino', 'vacunas',
        'situacion_laboral', 'vivienda_tenencia'
    ]

    for fig_key in orden_graficos:
        fig_obj = figs.get(fig_key) 
        
        if fig_obj is None: continue 
        
        if fig_key == 'piramide':
            document.add_paragraph('Gráficos Demográficos', style='Titulo2')
        elif fig_key == 'enfermedades':
            document.add_paragraph('Gráficos Socio-Sanitarios', style='Titulo2')
        elif fig_key == 'situacion_laboral':
            document.add_paragraph('Gráficos Socio-Económicos', style='Titulo2')
            
        img_stream = io.BytesIO()
        fig_obj.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        
        width = Inches(6.5) if fig_key == 'piramide' else Inches(6.0)
        try:
            document.add_picture(img_stream, width=width)
            if fig_obj.axes:
                p = document.add_paragraph(f"Gráfico: {fig_obj.axes[0].get_title()}", style='Caption')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if fig_key == 'enfermedades' and tasas_prevalencia_enf:
                p_enf = document.add_paragraph()
                p_enf.add_run('Tasas de Prevalencia (por 100 habitantes):').bold = True
                
                if not tasas_prevalencia_enf:
                    document.add_paragraph("No se calcularon tasas de prevalencia.", style='List Bullet')
                else:
                    for enfermedad, tasa in tasas_prevalencia_enf.items():
                        document.add_paragraph(f"{enfermedad}: ", style='List Bullet').add_run(f"{tasa} %").bold = True
                
                total_pob_unica = indices.get('Personas censadas (únicas)', 'N/A')
                p_enf_caption = document.add_paragraph(
                    f"(Calculado sobre una población única de {total_pob_unica} personas para el período '{anio}')"
                )
                p_enf_caption.runs[0].font.size = Pt(10)
                p_enf_caption.runs[0].italic = True

            if fig_key == 'enfermedades' and stats_enf_os:
                document.add_paragraph('Cobertura de Obra Social en Pacientes con Enfermedades Crónicas', style='Titulo2')
                
                table = document.add_table(rows=1, cols=6) 
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Enfermedad'
                hdr_cells[1].text = 'Total Pacientes'
                hdr_cells[2].text = 'Con Obra Social'
                hdr_cells[3].text = 'Sin Obra Social'
                hdr_cells[4].text = 'Con OS (%)'   
                hdr_cells[5].text = 'Sin OS (%)'   
                
                for enfermedad, data in stats_enf_os.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(enfermedad)
                    row_cells[1].text = str(data.get('Total Enfermos', 'N/A'))
                    row_cells[2].text = str(data.get('Con Obra Social', 'N/A'))
                    row_cells[3].text = str(data.get('Sin Obra Social', 'N/A'))
                    row_cells[4].text = str(data.get('Con OS (%)', 'N/A')) + ' %' 
                    row_cells[5].text = str(data.get('Sin OS (%)', 'N/A')) + ' %' 
                
                document.add_paragraph() 

        except Exception as e:
            print(f"Advertencia: No se pudo añadir la imagen {fig_key} al DOCX: {e}")

    # --- Sección 3: Datos y Proyecciones ---
    document.add_page_break()
    document.add_paragraph('Datos Detallados y Proyecciones', style='Titulo1')
    
    if df_piramide_data is not None and not df_piramide_data.empty:
        document.add_paragraph('Datos de la Pirámide Poblacional por Rango de Edad Quinquenal y Sexo', style='Titulo2')
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Grupo de Edad'
        hdr_cells[1].text = 'Hombres'
        hdr_cells[2].text = 'Mujeres'
        for index, row in df_piramide_data.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row["Grupo de Edad"])
            row_cells[1].text = str(row["Hombres"])
            row_cells[2].text = str(row["Mujeres"])

    document.add_paragraph('Interpretación y Proyecciones', style='Titulo2')
    p = document.add_paragraph()
    p.add_run('Forma de la Pirámide').bold = True
    p_interpretacion = document.add_paragraph()
    p_interpretacion.add_run(f"{interpretacion[0]}: ").bold = True
    p_interpretacion.add_run(interpretacion[1])

    p = document.add_paragraph()
    p.add_run('Proyección de Crecimiento Poblacional').bold = True
    for periodo, valor in proyecciones.items():
        document.add_paragraph(f"Población estimada en {periodo}: ", style='List Bullet').add_run(f"≈ {valor:,}").bold = True

    agregar_seccion_formulas(document)

    document.add_page_break() 
    document.add_paragraph('Origen de los datos', style='Titulo2')
    
    p_origen_1 = document.add_paragraph(
        "La información y estadísticas contenidas en este informe provienen de las cargas realizadas por el "
        "personal del CAPS u HOSPITAL y del archivo Excel descargado de dicho establecimiento."
    )
    
    p_origen_2 = document.add_paragraph()
    p_origen_2.add_run(
        "Por lo tanto, los resultados dependen directamente de la calidad, coherencia y "
        "actualización de los datos ingresados."
    ).italic = True

    # --- Guardar Documento ---
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# ---------------------------
# App principal
# ---------------------------
def main():
    st.set_page_config(page_title="Análisis Demográfico y Sanitario", layout="wide")
    
    # ------------------------------
    # LOGO Y ENCABEZADO DE BARRA LATERAL
    # ------------------------------
    with st.sidebar:
        st.header("🏥 Gestión Sanitaria")
        st.info("Espacio reservado para Logo Institucional o Publicidad del Ministerio/Hospital.")
        st.markdown("---")

    st.title("🔬 Análisis Demográfico y Socio-Sanitario Avanzado")
    st.subheader("Herramienta de análisis poblacional, caracterización de datos y estadísticas de salud")
    
    # ------------------------------
    # NUEVO: ADVERTENCIA IMPORTANTE (Antes de subir archivo)
    # ------------------------------
    st.warning("⚠️ **ATENCIÓN IMPORTANTE ANTES DE SUBIR EL ARCHIVO:**\n\n"
               "1. Asegúrese de **BORRAR** las columnas 'Latitud' y 'Longitud' de su Excel.\n"
               "2. Verifique y **ELIMINE** los DNI duplicados para evitar errores en el conteo.\n"
               "3. El archivo debe estar limpio de filas vacías innecesarias.")
    
    st.markdown("---")

    uploaded_file = st.file_uploader("📂 Subir archivo Excel", type=["xlsx"])

    if uploaded_file is not None:
        try:
            df_original = pd.read_excel(uploaded_file)
        except Exception as e:
            st.error(f"Error al leer el archivo Excel: {e}")
            return
        
        # --- VALIDACIÓN Y ESTANDARIZACIÓN DE COLUMNAS ---
        columnas_requeridas_original = ["DNI", "EDAD", "GENERO", "VIVIENDA_DIRECCION"]
        columnas_faltantes = [col for col in columnas_requeridas_original if col not in df_original.columns]
        
        if columnas_faltantes:
            st.error(f"⚠️ El archivo Excel debe contener, como mínimo, las siguientes columnas: {', '.join(columnas_faltantes)}.")
            st.info("Para las estadísticas de salud, se buscarán las columnas: `DIABETES`, `HTA`, `CHAGAS`, `TBC`, `CANCER`, `NIÑO_CONTROLES_COMPLETOS`, `NIÑO_VACUNACION_COMPLETA` y `OBRA_SOCIAL`.")
            return
        else:
            # Estandarizar el nombre de la columna de dirección para uso interno
            df_original.rename(columns={'VIVIENDA_DIRECCION': 'DIRECCION'}, inplace=True)
            
            # --- PARÁMETROS GLOBALES EN EL SIDEBAR ---
            # Solo se muestran si hay archivo cargado
            with st.sidebar:
                st.title("Parámetros Globales")
                localidad = st.sidebar.text_input("📍 Ingrese el nombre de la localidad", "Mi Localidad")
                nombre_informe_word = st.sidebar.text_input("📝 Ingrese el nombre para el Informe Word", "Informe Demográfico y Sanitario")
                
                st.sidebar.markdown("---")
                st.sidebar.subheader("Tasas de Simulación (Global)")
                natalidad_bruta = st.sidebar.number_input("Tasa de natalidad (x 1000)", value=15.0, step=0.1, format="%.1f")
                mortalidad_bruta = st.sidebar.number_input("Tasa de mortalidad (x 1000)", value=7.5, step=0.1, format="%.1f")
                tasa_migracion_neta = st.sidebar.number_input("Tasa de migración neta (x 1000)", value=1.0, step=0.1, format="%.1f")
                tfg = st.sidebar.number_input("Tasa Global Fecundidad (hijos/mujer)", value=2.1, step=0.1, format="%.1f")

            # --- SELECCIÓN DE COLUMNA DE FECHA Y EXTRACCIÓN DE AÑO ---
            st.markdown("### 1. Configuración de Fecha")
            columnas_fecha_posibles = [col for col in df_original.columns if 'FECHA' in col.upper()]
            default_index = 0
            if columnas_fecha_posibles:
                try:
                    default_index = list(df_original.columns).index(columnas_fecha_posibles[0])
                except ValueError:
                    pass 

            columna_fecha = st.selectbox(
                "Seleccione la columna que contiene la Fecha (para agrupar por año):", 
                df_original.columns,
                index=default_index
            )

            if not columna_fecha:
                st.info("Por favor, seleccione una columna de fecha para continuar.")
                return

            # --- DIAGNÓSTICO DE DATOS ---
            st.markdown("### 2. Diagnóstico de Datos")
            
            total_filas_inicial = len(df_original)
            st.info(f"Total de filas leídas del Excel: **{total_filas_inicial}**")

            try:
                # 1. Intentar convertir las fechas
                # Añadimos dayfirst=True para que entienda "DD/MM/YYYY"
                fechas_convertidas = pd.to_datetime(
                    df_original[columna_fecha], 
                    errors='coerce', 
                    dayfirst=True
                )
                
                # 2. Contar las filas que fallaron (NaT)
                filas_con_fecha_invalida = fechas_convertidas.isna()
                conteo_fechas_invalidas = filas_con_fecha_invalida.sum()
                
                with st.expander("Ver Diagnóstico de la Columna de Fecha"):
                    st.write(f"Columna de fecha seleccionada: `{columna_fecha}`")
                    st.warning(f"Filas con fecha inválida (vacías o texto): **{conteo_fechas_invalidas}**")
                    st.success(f"Filas con fecha válida: **{total_filas_inicial - conteo_fechas_invalidas}**")
                    st.markdown("---")
                    
                    if conteo_fechas_invalidas > 0:
                        st.error(f"El programa **descartará** las filas con fechas inválidas. Aquí hay 20 ejemplos de los datos en `{columna_fecha}` que están causando el problema:")
                        st.dataframe(df_original[filas_con_fecha_invalida][columna_fecha].head(20))
                    else:
                        st.success("¡Excelente! Todas las fechas son válidas.")

                # --- FIN DEL DIAGNÓSTICO ---
                
                # --- APLICAR EL FILTRO REAL ---
                df_original['ANIO'] = fechas_convertidas.dt.year
                df_original.dropna(subset=['ANIO'], inplace=True) # Este es el filtro
                df_original['ANIO'] = df_original['ANIO'].astype(int)
                
                anios_disponibles = sorted(df_original['ANIO'].unique())
                
                if not anios_disponibles:
                    st.error("No se encontraron años válidos en la columna de fecha seleccionada. Verifique los datos.")
                    return
                    
            except Exception as e:
                st.error(f"Error al procesar la columna de fecha '{columna_fecha}': {e}")
                return

            st.markdown("### 3. Vista Previa de Datos (Filtrados)")
            st.markdown(f"👀 Vista previa de los datos **después** de eliminar filas sin fecha válida. Quedan **{len(df_original)}** filas.")
            st.dataframe(df_original.head())
            
            # --- MOSTRAR EXPLICACIÓN DE FÓRMULAS ---
            mostrar_explicacion_formulas()

            st.markdown("---")

            # --- CREACIÓN DE PESTAÑAS (TABS) POR AÑO ---
            st.markdown(f"### 4. Análisis por Año para: {localidad.title()}")
            
            tab_nombres = ["Todos los Años"] + [str(a) for a in anios_disponibles]
            tabs = st.tabs(tab_nombres)

            # --- BUCLE DE ANÁLISIS POR CADA PESTAÑA ---
            for i, tab in enumerate(tabs):
                with tab:
                    if i == 0: # Pestaña "Todos los Años"
                        anio_seleccionado = "Todos"
                        df_original_filtrado = df_original
                    else: # Pestaña de un año específico
                        anio_seleccionado = anios_disponibles[i-1]
                        df_original_filtrado = df_original[df_original['ANIO'] == anio_seleccionado]
                    
                    st.header(f"Análisis del Período: {anio_seleccionado}")
                    st.caption(f"Usando los parámetros globales de la barra lateral.")

                    # --- PROCESAMIENTO DE DATOS (Filtrado) ---
                    df_original_filtrado = df_original_filtrado.copy()
                    df_original_filtrado.dropna(subset=['DNI'], inplace=True)
                    df_sin_duplicados_filtrado = df_original_filtrado.drop_duplicates(subset=['DNI'], keep='first').copy()
                    
                    if df_sin_duplicados_filtrado.empty:
                        st.warning(f"No hay datos de personas (con DNI único) para el período '{anio_seleccionado}'.")
                        continue # Saltar al siguiente tab

                    df_sin_duplicados_filtrado["PEA_estimado"] = df_sin_duplicados_filtrado.apply(lambda row: estimar_pea(row["EDAD"], row["GENERO"]), axis=1)

                    # --- CÁLCULO DE ÍNDICES ---
                    indices = calcular_indices_avanzados(df_original_filtrado, df_sin_duplicados_filtrado)
                    
                    st.subheader(f"📊 Resultados del Censo ({anio_seleccionado})")
                    col1, col2 = st.columns(2)
                    col1.metric("👥 Personas Censadas (Únicas)", f"{indices['Personas censadas (únicas)']}")
                    col2.metric("🏠 Viviendas Relevadas (Estimado)", f"{indices['Viviendas relevadas (estimado)']}")
                    st.info("La estimación de viviendas se basa en el conteo de direcciones únicas ('VIVIENDA_DIRECCION') para el período seleccionado.")

                    # --- ANÁLISIS SOCIO-ECONÓMICO (EN EXPANDER) ---
                    with st.expander(f"🔍 Análisis Socio-Económico Detallado ({anio_seleccionado})"):
                        st.markdown("#### Estadísticas de Variables Numéricas")
                        st.dataframe(df_sin_duplicados_filtrado.describe(include=np.number))
                        st.markdown("#### Estadísticas de Variables Categóricas (Texto)")
                        st.dataframe(df_sin_duplicados_filtrado.describe(include=['object']))
                        
                        st.markdown("---")
                        st.markdown("#### Gráficos de Variables Socio-Económicas Clave")
                        
                        figs_adicionales = graficar_analisis_adicional(df_sin_duplicados_filtrado)
                        if figs_adicionales:
                            for key, fig in figs_adicionales.items():
                                st.pyplot(fig)
                        else:
                            st.write("No se encontraron columnas socio-económicas (Ej: SITUACION_LABORAL, VIVIENDA_TENENCIA) o no hay datos para ellas en este período.")

                    # --- NUEVO: ANÁLISIS SOCIO-SANITARIO ---
                    st.markdown("---")
                    st.subheader(f"⚕️ Estadísticas Socio-Sanitarias ({anio_seleccionado})")
                    
                    N_total_poblacion = indices['Personas censadas (únicas)']
                    
                    stats_salud, filtros_salud = calcular_estadisticas_salud(df_sin_duplicados_filtrado.copy()) 
                    
                    figs_salud = graficar_estadisticas_salud(stats_salud, localidad, N_total_poblacion)

                    # --- Extraer datos de OS ---
                    stats_os_general = stats_salud.get('obra_social_general', {})
                    stats_enf_os = stats_salud.get('enfermedades_obra_social', {})
                    stats_embarazadas = stats_salud.get('embarazadas', {}) 
                    
                    stats_mayores_65_os = stats_salud.get('mayores_65_os', {})
                    stats_menores_15_os = stats_salud.get('menores_15_os', {})
                    

                    # --- Mostrar Métrica de Cobertura General ---
                    if stats_os_general:
                        st.metric(
                            label="Índice de Mutualización General (con OS)", 
                            value=f"{stats_os_general.get('Indice Mutualizacion (%)', 0)}%",
                            help=f"Del total de {stats_os_general.get('Población Total (única)', 0)} personas, {stats_os_general.get('Con Obra Social', 0)} tienen OS."
                        )
                    else:
                        st.info("No se encontró la columna `OBRA_SOCIAL` para calcular el índice de mutualización.")
                    
                    # --- Mostrar Métrica de Embarazadas con OS ---
                    if stats_embarazadas:
                        total_emb = stats_embarazadas.get('Total Estimado', 0)
                        con_os_emb = stats_embarazadas.get('Con Obra Social', 0)
                        sin_os_emb = stats_embarazadas.get('Sin Obra Social', 0)
                        st.metric(
                            label="🤰 Embarazadas (Estimado)", 
                            value=f"{total_emb}",
                            help=f"Total: {total_emb} | Con OS: {con_os_emb} | Sin OS: {sin_os_emb}. (Cálculo basado en MUJER_FUM y fecha actual)" 
                        )
                    else:
                        st.info("No se encontró la columna `MUJER_FUM` para estimar embarazos.")

                    # --- Mostrar Métricas de Edad y OS ---
                    if stats_mayores_65_os:
                        total_mayores = stats_mayores_65_os.get('Total Mayores 65', 0)
                        con_os_mayores = stats_mayores_65_os.get('Con Obra Social', 0)
                        sin_os_mayores = stats_mayores_65_os.get('Sin Obra Social', 0)
                        st.metric(
                            label="👴 Población 65+ años",
                            value=f"{total_mayores}",
                            help=f"Total: {total_mayores} | Con OS: {con_os_mayores} | Sin OS: {sin_os_mayores}"
                        )

                    if stats_menores_15_os:
                        total_menores = stats_menores_15_os.get('Total Menores 15 (0-14)', 0)
                        con_os_menores = stats_menores_15_os.get('Con Obra Social', 0)
                        sin_os_menores = stats_menores_15_os.get('Sin Obra Social', 0)
                        st.metric(
                            label="👶 Población 0-14 años",
                            value=f"{total_menores}",
                            help=f"Total: {total_menores} | Con OS: {con_os_menores} | Sin OS: {sin_os_menores}"
                        )

                    st.markdown("---") 

                    # --- Expander para Cobertura en Enfermedades ---
                    with st.expander(f"🔍 Ver Cobertura de Obra Social en Enfermedades Crónicas ({anio_seleccionado})"):
                        if stats_enf_os:
                            df_enf_os = pd.DataFrame.from_dict(stats_enf_os, orient='index')
                            st.markdown("#### Cobertura de OS por Enfermedad")
                            st.dataframe(df_enf_os)
                            st.caption("Esta tabla muestra, del total de pacientes con una enfermedad (incluyendo DUAL), cuántos tienen OS y el porcentaje que representan.")
                        else:
                            st.info("No se encontraron datos de 'OBRA_SOCIAL' o no hay casos de enfermedades para analizar la cobertura.")


                    # --- Cálculo de Tasa de Prevalencia de Enfermedades ---
                    tasas_prevalencia = {}
                    total_poblacion_unica = indices['Personas censadas (únicas)']
                    
                    if total_poblacion_unica > 0 and stats_salud['enfermedades']:
                        for enfermedad, casos in stats_salud['enfermedades'].items():
                            # Tasa por 100 habitantes (Prevalencia)
                            tasa = (casos / total_poblacion_unica) * 100
                            tasas_prevalencia[enfermedad] = round(tasa, 2)

                    col_salud1, col_salud2 = st.columns(2)
                    
                    with col_salud1:
                        st.markdown("#### Prevalencia de Enfermedades")
                        if figs_salud.get('enfermedades'):
                            st.pyplot(figs_salud['enfermedades'])
                            
                            df_enf_stats = pd.DataFrame.from_dict(
                                stats_salud['enfermedades'], 
                                orient='index', 
                                columns=['Casos Positivos']
                            )
                            
                            if tasas_prevalencia:
                                df_enf_stats['Prevalencia (%)'] = df_enf_stats.index.map(tasas_prevalencia)
                                df_enf_stats['Prevalencia (%)'] = df_enf_stats['Prevalencia (%)'].fillna(0)
                            
                            st.dataframe(df_enf_stats)
                            st.caption(f"La prevalencia (%) se calcula sobre el total de personas únicas ({total_poblacion_unica}).")

                        else:
                            st.info("No se encontraron datos o columnas de enfermedades (ej: DIABETES, HTA).")

                    with col_salud2:
                        st.markdown("#### Estado de Control Niño")
                        if figs_salud.get('control_nino'):
                            st.pyplot(figs_salud['control_nino'])
                            st.dataframe(pd.DataFrame.from_dict(stats_salud['control_nino'], orient='index', columns=['Cantidad']))
                        else:
                            st.info("No se encontró la columna 'NIÑO_CONTROLES_COMPLETOS' o no hay datos.")
                        
                        st.markdown("#### Estado de Vacunación")
                        if figs_salud.get('vacunas'):
                            st.pyplot(figs_salud['vacunas'])
                            st.dataframe(pd.DataFrame.from_dict(stats_salud['vacunas'], orient='index', columns=['Cantidad']))
                        else:
                            st.info("No se encontró la columna 'NIÑO_VACUNACION_COMPLETA' o no hay datos.")


                    # --- TASAS Y PROYECCIONES ---
                    st.markdown("---")
                    st.subheader(f"🔮 Interpretación y Proyecciones ({anio_seleccionado})")
                    tasas = calcular_tasas(indices, natalidad_bruta, mortalidad_bruta, tfg, tasa_migracion_neta)
                    st.metric("Tasa de Crecimiento Anual Estimada (%)", f'{tasas["Tasa de crecimiento anual (%)"]}%')
                    
                    interpretacion_piramide = interpretar_piramide(indices)
                    st.info(f"**Forma de la Pirámide: {interpretacion_piramide[0]}**")
                    st.write(interpretacion_piramide[1])

                    proyecciones = generar_proyeccion(indices, tasas)
                    st.subheader("Proyección de Crecimiento Poblacional (simulada)")
                    col1, col2, col3 = st.columns(3)
                    col1.metric("Población en 5 años", f"≈ {proyecciones['5 años']:,}")
                    col2.metric("Población en 10 años", f"≈ {proyecciones['10 años']:,}")
                    col3.metric("Población en 20 años", f"≈ {proyecciones['20 años']:,}")
                    st.markdown("---")

                    # --- GRÁFICOS DEMOGRÁFICOS PRINCIPALES ---
                    st.subheader(f"Gráficos Demográficos Principales ({anio_seleccionado})")
                    
                    fig_piramide, df_piramide_data = graficar_piramide(df_sin_duplicados_filtrado.copy(), localidad)
                    fig_genero = graficar_genero(df_sin_duplicados_filtrado, localidad)
                    
                    col_graf1, col_graf2 = st.columns([2, 1])
                    with col_graf1:
                        st.pyplot(fig_piramide)
                    with col_graf2:
                        st.pyplot(fig_genero)

                    st.markdown("---")
                    st.subheader(f"Tabla de la Pirámide Poblacional ({anio_seleccionado})")
                    st.dataframe(df_piramide_data.set_index("Grupo de Edad"))


                    # --------------------------------------------------------------------------------
                    # ANÁLISIS POR AGENTE SANITARIO
                    # --------------------------------------------------------------------------------
                    st.markdown("---")
                    st.subheader(f"🕵️ Análisis Individual por Agente Sanitario ({anio_seleccionado})")
                    
                    col_agente = None
                    if 'USUARIO_AGENTE_SANITARIO' in df_sin_duplicados_filtrado.columns:
                        col_agente = 'USUARIO_AGENTE_SANITARIO'
                    elif 'NOMBRE_AGENTE_SANITARIO' in df_sin_duplicados_filtrado.columns:
                        col_agente = 'NOMBRE_AGENTE_SANITARIO'
                    
                    if col_agente:
                        lista_agentes = sorted(df_sin_duplicados_filtrado[col_agente].dropna().astype(str).unique())
                        
                        if len(lista_agentes) > 0:
                            agente_seleccionado = st.selectbox(
                                f"Seleccione un Agente Sanitario para filtrar (Columna: {col_agente}):", 
                                lista_agentes,
                                key=f"select_agente_{anio_seleccionado}"
                            )
                            
                            if agente_seleccionado:
                                # --- 1. FILTRAR DATOS ---
                                df_agente = df_sin_duplicados_filtrado[df_sin_duplicados_filtrado[col_agente].astype(str) == agente_seleccionado].copy()
                                
                                st.markdown(f"### Resultados para: **{agente_seleccionado}**")
                                st.info(f"Mostrando estadísticas exclusivas para la población a cargo de este agente.")

                                # --- 2. CALCULAR MÉTRICAS AGENTE ---
                                indices_agente = calcular_indices_avanzados(df_agente, df_agente)
                                stats_salud_agente, _ = calcular_estadisticas_salud(df_agente) # No necesitamos filtros booleanos aquí
                                N_agente = indices_agente['Personas censadas (únicas)']

                                # --- 3. MOSTRAR MÉTRICAS PRINCIPALES ---
                                col_ag1, col_ag2, col_ag3 = st.columns(3)
                                col_ag1.metric("👥 Personas (Agente)", f"{N_agente}")
                                col_ag2.metric("🏠 Viviendas (Estimado)", f"{indices_agente['Viviendas relevadas (estimado)']}")
                                
                                os_gral_agente = stats_salud_agente.get('obra_social_general', {})
                                val_os_agente = os_gral_agente.get('Indice Mutualizacion (%)', 0)
                                col_ag3.metric("🏥 Cobertura OS General", f"{val_os_agente}%")

                                # --- 4. GRÁFICOS AGENTE ---
                                st.markdown("#### 📊 Pirámide y Género (Agente)")
                                fig_pyr_ag, df_piramide_agente = graficar_piramide(df_agente, f"Agente: {agente_seleccionado}")
                                fig_gen_ag = graficar_genero(df_agente, f"Agente: {agente_seleccionado}")

                                col_ga1, col_ga2 = st.columns([2, 1])
                                with col_ga1:
                                    st.pyplot(fig_pyr_ag)
                                with col_ga2:
                                    st.pyplot(fig_gen_ag)
                                
                                st.markdown("#### ⚕️ Salud y Enfermedades (Agente)")
                                figs_salud_agente = graficar_estadisticas_salud(stats_salud_agente, f"Agente: {agente_seleccionado}", N_agente)
                                
                                if figs_salud_agente.get('enfermedades'):
                                    st.pyplot(figs_salud_agente['enfermedades'])
                                else:
                                    st.write("Sin datos de enfermedades para este agente.")
                                
                                col_vac_ag1, col_vac_ag2 = st.columns(2)
                                with col_vac_ag1:
                                    if figs_salud_agente.get('control_nino'):
                                        st.pyplot(figs_salud_agente['control_nino'])
                                    else:
                                        st.caption("Sin datos de control niño.")
                                with col_vac_ag2:
                                    if figs_salud_agente.get('vacunas'):
                                        st.pyplot(figs_salud_agente['vacunas'])
                                    else:
                                        st.caption("Sin datos de vacunas.")

                                # --- 5. DETALLE ENFERMOS AGENTE ---
                                with st.expander(f"Ver detalle numérico de Salud ({agente_seleccionado})"):
                                    col_det1, col_det2 = st.columns(2)
                                    with col_det1:
                                        st.write("**Enfermedades:**")
                                        st.write(stats_salud_agente.get('enfermedades', "Sin casos"))
                                    with col_det2:
                                        st.write("**Cobertura OS en Enfermos:**")
                                        st.write(stats_salud_agente.get('enfermedades_obra_social', "Sin datos"))

                                # --- 6. DESCARGA EXCEL POR AGENTE ---
                                st.markdown("---")
                                st.markdown("##### 📥 Descargar Estadísticas de este Agente")
                                
                                excel_buffer_agente = io.BytesIO()
                                with pd.ExcelWriter(excel_buffer_agente, engine='openpyxl') as writer:
                                    resumen_data = {
                                        'Indicador': ['Personas (Agente)', 'Viviendas (Estimado)', 'Cobertura OS General (%)'],
                                        'Valor': [
                                            N_agente, 
                                            indices_agente['Viviendas relevadas (estimado)'], 
                                            f"{val_os_agente}%"
                                        ]
                                    }
                                    pd.DataFrame(resumen_data).to_excel(writer, index=False, sheet_name='Resumen_General')

                                    if not df_piramide_agente.empty:
                                        df_piramide_agente.to_excel(writer, index=False, sheet_name='Piramide_Poblacional')
                                    
                                    if stats_salud_agente['enfermedades']:
                                        df_enf_stats_agente = pd.DataFrame.from_dict(
                                            stats_salud_agente['enfermedades'], 
                                            orient='index', 
                                            columns=['Casos Positivos']
                                        )
                                        if N_agente > 0:
                                            df_enf_stats_agente['Prevalencia (%)'] = (df_enf_stats_agente['Casos Positivos'] / N_agente) * 100
                                            df_enf_stats_agente['Prevalencia (%)'] = df_enf_stats_agente['Prevalencia (%)'].round(2)
                                        df_enf_stats_agente.to_excel(writer, sheet_name='Enfermedades')
                                    
                                    stats_enf_os_agente = stats_salud_agente.get('enfermedades_obra_social', {})
                                    if stats_enf_os_agente:
                                        pd.DataFrame.from_dict(stats_enf_os_agente, orient='index').to_excel(writer, sheet_name='Cobertura_OS_Enfermedad')
                                    
                                    stats_os_gen_agente = stats_salud_agente.get('obra_social_general', {})
                                    if stats_os_gen_agente:
                                        pd.DataFrame.from_dict(stats_os_gen_agente, orient='index', columns=['Valor']).to_excel(writer, sheet_name='Cobertura_OS_General')
                                    
                                    stats_emb_agente = stats_salud_agente.get('embarazadas', {})
                                    if stats_emb_agente:
                                        pd.DataFrame.from_dict(stats_emb_agente, orient='index', columns=['Valor']).to_excel(writer, sheet_name='Embarazadas_Estimado')

                                    stats_may65_agente = stats_salud_agente.get('mayores_65_os', {})
                                    if stats_may65_agente:
                                        pd.DataFrame.from_dict(stats_may65_agente, orient='index', columns=['Valor']).to_excel(writer, sheet_name='Mayores_65_OS')
                                    
                                    stats_men15_agente = stats_salud_agente.get('menores_15_os', {})
                                    if stats_men15_agente:
                                        pd.DataFrame.from_dict(stats_men15_agente, orient='index', columns=['Valor']).to_excel(writer, sheet_name='Menores_15_OS')

                                    if stats_salud_agente['control_nino']:
                                        pd.DataFrame.from_dict(stats_salud_agente['control_nino'], orient='index', columns=['Cantidad']).to_excel(writer, sheet_name='Control_Nino')
                                    if stats_salud_agente['vacunas']:
                                        pd.DataFrame.from_dict(stats_salud_agente['vacunas'], orient='index', columns=['Cantidad']).to_excel(writer, sheet_name='Vacunas')

                                excel_buffer_agente.seek(0)
                                st.download_button(
                                    label=f"⬇️ Descargar Excel de {agente_seleccionado}",
                                    data=excel_buffer_agente,
                                    file_name=f"Estadisticas_{agente_seleccionado.replace(' ', '_')}_{anio_seleccionado}.xlsx",
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                    key=f"download_excel_agente_{agente_seleccionado}_{anio_seleccionado}"
                                )

                        else:
                            st.warning(f"La columna `{col_agente}` existe pero no contiene datos válidos.")
                    else:
                        st.info("No se encontró la columna `USUARIO_AGENTE_SANITARIO` (ni `NOMBRE_AGENTE_SANITARIO`) para realizar el desglose por agente.")

                    # --- DESCARGAS ---
                    st.markdown("---")
                    st.subheader("📥 Zona de Descargas")
                    
                    excel_buffer = io.BytesIO()
                    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                        df_piramide_data.to_excel(writer, index=False, sheet_name='Piramide_Poblacional')
                        
                        if stats_salud['enfermedades']:
                            df_enf_stats_excel = pd.DataFrame.from_dict(
                                stats_salud['enfermedades'], 
                                orient='index', 
                                columns=['Casos Positivos']
                            )
                            if tasas_prevalencia:
                                df_enf_stats_excel['Prevalencia (%)'] = df_enf_stats_excel.index.map(tasas_prevalencia)
                                df_enf_stats_excel['Prevalencia (%)'] = df_enf_stats_excel['Prevalencia (%)'].fillna(0)
                            df_enf_stats_excel.to_excel(writer, sheet_name='Enfermedades')

                        if stats_enf_os:
                            pd.DataFrame.from_dict(stats_enf_os, orient='index').to_excel(writer, sheet_name='Cobertura_OS_Enfermedad')
                        if stats_os_general:
                            pd.DataFrame.from_dict(stats_os_general, orient='index', columns=['Valor']).to_excel(writer, sheet_name='Cobertura_OS_General')
                        
                        if stats_embarazadas:
                            pd.DataFrame.from_dict(stats_embarazadas, orient='index', columns=['Valor']).to_excel(writer, sheet_name='Embarazadas_Estimado')
                        
                        if stats_mayores_65_os:
                            pd.DataFrame.from_dict(stats_mayores_65_os, orient='index', columns=['Valor']).to_excel(writer, sheet_name='Mayores_65_OS')
                        if stats_menores_15_os:
                            pd.DataFrame.from_dict(stats_menores_15_os, orient='index', columns=['Valor']).to_excel(writer, sheet_name='Menores_15_OS')

                        if stats_salud['control_nino']:
                                pd.DataFrame.from_dict(stats_salud['control_nino'], orient='index', columns=['Cantidad']).to_excel(writer, sheet_name='Control_Nino')
                        if stats_salud['vacunas']:
                                pd.DataFrame.from_dict(stats_salud['vacunas'], orient='index', columns=['Cantidad']).to_excel(writer, sheet_name='Vacunas')

                    excel_buffer.seek(0)
                    st.download_button(
                        label="Descargar datos (Pirámide y Salud) en Excel",
                        data=excel_buffer,
                        file_name=f"Datos_Demograficos_Salud_{localidad.replace(' ', '_')}_{anio_seleccionado}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key=f"download_excel_{anio_seleccionado}" 
                    )

                    st.markdown("---")
                    
                    st.subheader(f"📥 Descargar Informe Completo en Word ({anio_seleccionado})")
                    
                    if st.button(f"Generar Informe para {anio_seleccionado}", key=f"generate_word_{anio_seleccionado}"):
                        if nombre_informe_word: 
                            with st.spinner(f"Generando informe para {anio_seleccionado}..."):
                                figs_dict = {
                                    'piramide': fig_piramide, 
                                    'genero': fig_genero
                                }
                                figs_dict.update(figs_adicionales)
                                figs_dict.update(figs_salud)

                                informe_word_buffer = generar_informe_word(
                                    indices, tasas, proyecciones, 
                                    interpretacion_piramide, 
                                    nombre_informe_word, 
                                    figs_dict, 
                                    df_piramide_data,
                                    anio=anio_seleccionado, 
                                    tasas_prevalencia_enf=tasas_prevalencia,
                                    stats_os_general=stats_os_general, 
                                    stats_enf_os=stats_enf_os, 
                                    stats_embarazadas=stats_embarazadas, 
                                    stats_mayores_65_os=stats_mayores_65_os, 
                                    stats_menores_15_os=stats_menores_15_os  
                                )
                                
                                st.download_button(
                                    label="✅ ¡Informe Listo! Haga clic para descargar",
                                    data=informe_word_buffer,
                                    file_name=f"Informe_{nombre_informe_word.replace(' ', '_')}_{anio_seleccionado}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"download_word_button_{anio_seleccionado}" 
                                )
                        else:
                            st.warning("Por favor, ingrese un nombre para el informe en la barra lateral izquierda antes de generarlo.")

                    # --- SECCIÓN DESCARGAR PADRONES ---
                    st.markdown("---")
                    st.subheader(f"🧑‍🤝‍🧑 Descarga de Padrones Específicos ({anio_seleccionado})")
                    st.info(
                        "Genere un archivo Excel con los datos de las personas que cumplen con cada criterio de salud "
                        "(una solapa por criterio)."
                    )

                    df_para_padrones = df_sin_duplicados_filtrado.copy()
                    
                    if 'DIRECCION' in df_para_padrones.columns:
                        df_para_padrones.rename(columns={'DIRECCION': 'DOMICILIO'}, inplace=True)

                    columnas_deseadas = [
                        'DNI', 
                        'NOMBRE', 
                        'APELLIDO', 
                        'EDAD',               
                        'GENERO',             
                        'DOMICILIO',          
                        'NOMBRE_AGENTE_SANITARIO',
                        'USUARIO_AGENTE_SANITARIO' 
                    ]
                    
                    cols_padron_finales = [col for col in columnas_deseadas if col in df_para_padrones.columns]

                    padron_excel_buffer = io.BytesIO()
                    try:
                        with pd.ExcelWriter(padron_excel_buffer, engine='openpyxl') as writer:
                            
                            if not filtros_salud:
                                pd.DataFrame(["No se generaron filtros de salud."]).to_excel(writer, sheet_name='Error', index=False)
                            else:
                                for nombre_filtro, filtro_bool in filtros_salud.items():
                                    
                                    padron_df = df_para_padrones[filtro_bool]
                                    padron_df_final = padron_df[cols_padron_finales]
                                    
                                    sheet_name = nombre_filtro.replace(' (', '_').replace('+', '').replace(')', '').replace(' ', '_')
                                    sheet_name = sheet_name.replace('á', 'a').replace('é', 'e').replace('í', 'i').replace('ó', 'o').replace('ú', 'u')
                                    sheet_name = sheet_name[:31] 
                                    
                                    if not padron_df_final.empty:
                                        padron_df_final.to_excel(writer, sheet_name=sheet_name, index=False)
                                    else:
                                        pd.DataFrame([f"No hay personas que cumplan con el criterio '{nombre_filtro}'."]).to_excel(writer, sheet_name=sheet_name, index=False)
                        
                        padron_excel_buffer.seek(0)
                        
                        st.download_button(
                            label="⬇️ Descargar Padrón de Grupos de Riesgo (Excel)",
                            data=padron_excel_buffer,
                            file_name=f"Padron_Riesgo_{localidad.replace(' ', '_')}_{anio_seleccionado}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key=f"download_padron_{anio_seleccionado}" 
                        )

                    except Exception as e:
                        st.error(f"Error al generar el Excel del padrón: {e}")

    else:
        st.info("☝️ Sube un archivo Excel para comenzar el análisis.")
        st.markdown("""
        Asegúrate de que tu archivo contenga, como mínimo, las columnas: 
        `DNI`, `EDAD`, `GENERO`, `VIVIENDA_DIRECCION` y una columna de **Fecha**.
        
        Para las estadísticas de salud, la aplicación buscará (si existen):
        * `DIABETES`, `HTA`, `CHAGAS`, `TBC`, `CANCER` (con valores 'SI'/'NO')
        * `NIÑO_CONTROLES_COMPLETOS` (con valores 'Completo'/'Incompleto' o 'SI'/'NO')
        * `NIÑO_VACUNACION_COMPLETA` (con valores 'Completas'/'Incompletas' o 'SI'/'NO')
        * `OBRA_SOCIAL` (con valores 'SI'/'NO')
        * `MUJER_FUM` (como fecha, para estimar embarazos)
        * **`USUARIO_AGENTE_SANITARIO`** (para el desglose por agente)
        
        Para la descarga de padrones, se buscarán también (si existen):
        * `NOMBRE`, `APELLIDO`, `NOMBRE_AGENTE_SANITARIO`
        
        (Las categorías **DUAL (DBT+HTA)** y **Embarazadas** se calcularán automáticamente)
        """)
    
    # ------------------------------
    # PIE DE PÁGINA PRINCIPAL (Main Footer)
    # ------------------------------
    st.markdown("---") # Línea separadora
    col_footer_1, col_footer_2 = st.columns(2)
    
    with col_footer_1:
        st.markdown("**Diseñado y Desarrollado por:**")
        st.markdown("💻 Tec. Muñoz, Mauricio Ivan")
    
    with col_footer_2:
        st.caption("⚠️ **Descargo de Responsabilidad:**")
        st.caption(
            "Esta aplicación puede contener errores o inconsistencias. "
            "Los resultados dependen directamente de la calidad de los datos cargados. "
            "Se ruega corroborar toda la información antes de su presentación oficial."
        )

    # ------------------------------
    # PIE DE BARRA LATERAL (Contacto y Donación)
    # ------------------------------
    # Esta sección se agrega SIEMPRE al final de la barra lateral
    with st.sidebar:
        st.markdown("---")
        st.subheader("☕ Apoyo al Desarrollador")
        st.write("Si la aplicación te fue útil, puedes agradecer con una donación 👇")
        
        # Botón HTML personalizado para que se vea "ordenado" y rojo
        st.markdown("""
            <a href="https://tecito.app/ivannefa" target="_blank">
                <button style="
                    width: 100%;
                    background-color: #ef4444;
                    border: none;
                    color: white;
                    padding: 10px 20px;
                    text-align: center;
                    text-decoration: none;
                    display: inline-block;
                    font-size: 16px;
                    margin: 4px 2px;
                    cursor: pointer;
                    border-radius: 8px;
                    font-weight: bold;
                ">
                    ☕ Invitame un Tecito
                </button>
            </a>
            """, unsafe_allow_html=True)

        st.markdown("---")
        st.subheader("📧 Feedback & Ideas")
        st.write("¿Tienes ideas para mejorar la app o encontraste un error? Escríbeme:")
        
        # Botón mailto
        st.link_button("✉️ Enviar correo a Soporte", "mailto:hhepatalgina@gmail.com")


if __name__ == "__main__":
    main()
